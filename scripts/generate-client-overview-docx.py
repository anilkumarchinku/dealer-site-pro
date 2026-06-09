from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s]
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.font.name = 'Calibri'

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
    return table

def add_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('Dealer Site Pro', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Client-Facing Overview')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_line()

# --- What is Dealer Site Pro? ---
add_heading('What is Dealer Site Pro?')
add_para('Dealer Site Pro is a website builder platform designed exclusively for Indian automotive dealerships. It allows any dealer \u2014 whether they sell cars, bikes, or auto-rickshaws \u2014 to launch a fully functional, professional website without writing a single line of code.')

add_line()

# --- The Problem It Solves ---
add_heading('The Problem It Solves')
add_para('Most small and mid-size dealerships in India either have no website at all, or rely on outdated, expensive ones built by agencies. They lose customers to competitors who show up on Google. They manage leads on paper or WhatsApp. They have no online presence for their inventory.')
add_para('Dealer Site Pro fixes all of this in one platform.')

add_line()

# --- How It Works ---
add_heading('How It Works (Dealer Journey)')
add_numbered('Sign up ', 'Step 1: ')
add_numbered('Walk through a guided setup (5-6 steps) \u2014 enter dealership details, pick brands you sell, choose services you offer, select a design template, customize your homepage', 'Step 2: ')
add_numbered('Go live \u2014 get a free website instantly at yourname.indrav.in', 'Step 3: ')
add_numbered('Manage everything from a single dashboard \u2014 inventory, leads, reviews, bookings, analytics', 'Step 4: ')
add_para('')
add_para('That\u2019s it. No developers, no hosting bills, no maintenance headaches.')

add_line()

# --- What Dealers Get ---
add_heading('What Dealers Get')

add_heading('A. A Professional Public Website', level=2)
add_bullet('Vehicle catalog with search and filters (by brand, price, fuel type, body type)')
add_bullet('Individual vehicle pages with full specs, photos, and pricing')
add_bullet('Built-in tools for customers \u2014 EMI calculator, on-road price estimator, insurance quotes')
add_bullet('Contact form, test drive booking, service appointment booking')
add_bullet('Customer reviews section')
add_bullet('Fully mobile-responsive and SEO-optimized (shows up on Google)')
add_bullet('4 design themes to choose from \u2014 Sporty, Modern, Family, Luxury')
add_bullet('Official brand colors applied automatically (Maruti blue, Tata teal, etc.)')

add_heading('B. A Complete Dealer Dashboard', level=2)
add_bullet(' \u2014 Add vehicles manually, look up details via RC number, upload photos, save drafts, mark as sold/reserved', 'Inventory Management')
add_bullet(' \u2014 Every inquiry, test drive request, and quote gets tracked with priority levels (hot/warm/cold) and status (new/contacted/converted/lost)', 'Lead Management (CRM)')
add_bullet(' \u2014 Collect, moderate, and reply to customer reviews. Sync reviews from Google.', 'Reviews')
add_bullet(' \u2014 All contact form submissions in one inbox', 'Messages')
add_bullet(' \u2014 Manage test drive and service appointments', 'Bookings')
add_bullet(' \u2014 Customers can submit their vehicle for sale or exchange', 'Sell/Exchange Requests')
add_bullet(' \u2014 Create and manage special deals', 'Offers & Promotions')
add_bullet(' \u2014 Track page views, unique visitors, leads generated, traffic sources, and conversions', 'Analytics')
add_bullet(' \u2014 Send browser notifications to website visitors', 'Push Notifications')
add_bullet(' \u2014 Start with a free subdomain, upgrade to your own custom domain anytime', 'Domain Management')

add_heading('C. Vehicle Type Support', level=2)
add_bullet(' \u2014 New cars, used cars, certified pre-owned', '4-Wheelers (Cars)')
add_bullet(' \u2014 New, used, electric', '2-Wheelers (Bikes & Scooters)')
add_bullet(' \u2014 Passenger, cargo, electric', '3-Wheelers (Auto-rickshaws)')
add_bullet('Each category has its own dedicated pages, catalog data, and management tools')

add_heading('D. Pre-Loaded Vehicle Data', level=2)
add_bullet('Thousands of vehicle models already in the system with specs, images, colors, and pricing')
add_bullet('Dealers just select what they sell \u2014 no need to enter data from scratch')
add_bullet('Covers all major Indian brands (Maruti Suzuki, Hyundai, Tata, Mahindra, Honda, Hero, Bajaj, TVS, Piaggio, and many more)')

add_line()

# --- APIs & Integrations ---
add_heading('APIs & Integrations \u2014 What Powers the Platform')

add_heading('Vehicle Verification APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['RC Number Lookup (Surepass)', 'Enter a vehicle registration number \u2192 instantly get owner name, make/model, fuel type, engine/chassis number, insurance status, fitness expiry, and pending challans. Costs \u20b93 per lookup, tracked via credit system. Results cached for 24 hours.'],
        ['VIN Decoder (NHTSA)', 'Enter a VIN number \u2192 get make, model, year, fuel type decoded automatically'],
        ['Draft Vehicle Creation', 'RC lookup data auto-fills a draft vehicle listing. Dealer just adds photos and price to publish. Duplicate RC detection built in.'],
    ]
)

add_heading('Inventory Management APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Manual Vehicle CRUD', 'Add, edit, delete vehicles with full details \u2014 make, model, variant, year, price, mileage, color, condition, photos'],
        ['Image Upload', 'Upload vehicle photos to cloud storage with CDN delivery'],
        ['Cyepro DMS Sync', 'Import inventory directly from Cyepro dealer management system. Server-side proxy with pagination, filtering, and price mapping. Includes diagnostic/test endpoint.'],
        ['AI Description Generator', 'Uses Claude AI to auto-generate professional vehicle listing descriptions from specs. Rate-limited, input-sanitized.'],
        ['Used Vehicle Price Offers', 'Set custom pricing on used vehicles from any source (manual or Cyepro)'],
    ]
)

add_heading('Lead & CRM APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Lead Submission', 'Captures inquiries from dealer website. 5-minute duplicate detection. Auto-notifies dealer via SMS and email. Auto-forwards to Cyepro CRM if configured.'],
        ['Test Drive Booking', 'Customers book test drives with preferred date/time. Creates both lead and booking records.'],
        ['Lead Management', 'List, filter, and update lead status (new \u2192 contacted \u2192 qualified \u2192 converted \u2192 lost) with priority tracking'],
    ]
)

add_heading('Service & Booking APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Car Service Booking', 'Customers book service appointments \u2014 general service, body repair, AC service, etc. Supports home pickup. Rate-limited.'],
        ['2W/3W Service Booking', 'Dedicated service booking for bikes and autos'],
        ['Service Centers', 'Manage multiple service center locations with working hours, images, and pricing tiers'],
    ]
)

add_heading('Customer Engagement APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Review Submission', 'Customers submit reviews (rate-limited to 5/day per IP). Auto-approve or manual moderation.'],
        ['Review Moderation', 'Dealers approve/reject/flag reviews, reply to them, feature best ones on homepage'],
        ['Google Review Sync', 'Import reviews from Google Maps/Places automatically using Place ID'],
        ['Sell/Exchange Requests', 'Customers submit their vehicle for sale \u2014 with photos, expected price, preferred inspection slot. Auto-emails confirmation. When dealer approves, auto-creates inventory listing.'],
        ['Customer Panel', 'Customers look up their history (inquiries, test drives, sell requests) by phone/email'],
        ['Push Notifications', 'Subscribe visitors to web push. Dealers broadcast notifications (new arrivals, price drops, announcements).'],
    ]
)

add_heading('Payment & Subscription APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Razorpay Subscription', 'Create subscriptions for Pro/Premium plans. Handles trial periods.'],
        ['Payment Verification', 'Verify Razorpay payments with signature validation. Idempotent \u2014 prevents duplicate charges.'],
        ['Vehicle Booking Payments', 'Customers pay booking amounts online for 2W/3W vehicles. Order creation + payment verification.'],
        ['Razorpay Webhooks', 'Handles subscription lifecycle events \u2014 activated, charged, cancelled, payment failed. Deduplication built in.'],
    ]
)

add_heading('Domain Management APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Subdomain Creation', 'Auto-generates a free subdomain from business name + city'],
        ['Custom Domain Connect', 'Connect your own domain. Registers on Vercel. Returns DNS records to configure.'],
        ['DNS Verification', 'Verify A and CNAME records are configured correctly. Saves verification history.'],
        ['Domain Search', 'Search available domains for purchase (Premium tier)'],
        ['Domain Monitoring', 'Cron job checks SSL certificates and domain expiry every 12 hours'],
        ['Domain Resolution', 'Edge middleware resolves subdomain/custom domain to dealer slug with Redis caching'],
    ]
)

add_heading('Marketplace & Social APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Cross-Dealer Marketplace', 'Search across all dealers\u2019 inventory \u2014 filter by make, fuel, condition, body type, price, location'],
        ['Social Media Auto-Post', 'Auto-post vehicle listings to Facebook Page, Instagram Business, and Twitter/X'],
    ]
)

add_heading('Auth & Account APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['OTP Login', 'Send OTP to email for passwordless authentication. Rate-limited.'],
        ['Registration Check', 'Verify email/phone availability before signup'],
        ['Account Deletion', 'GDPR-compliant account deletion with PII anonymization'],
    ]
)

add_heading('Admin & System APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Admin Dashboard', 'List all dealers, manage catalog, deploy templates'],
        ['Template Deployment', 'Push template + brand updates to dealer sites with ISR revalidation'],
        ['Health Check', 'Monitor Supabase, Razorpay, Vercel connectivity'],
        ['API Credit Tracking', 'Track per-dealer API usage and costs (RC lookups, etc.)'],
    ]
)

add_heading('Utility APIs', level=2)
add_table(
    ['API', 'What It Does'],
    [
        ['Finance Pre-Check', 'Redirect to finance partner for loan eligibility'],
        ['FASTag Recharge', 'Redirect to FASTag recharge partner'],
        ['Brand Catalog', 'Get models for specific make with preview data'],
    ]
)

add_line()

# --- External Service Integrations ---
add_heading('External Service Integrations')
add_table(
    ['Service', 'Purpose'],
    [
        ['Supabase', 'Database (PostgreSQL), Auth, Storage, Edge Functions'],
        ['Razorpay', 'Payments, subscriptions, webhooks'],
        ['Surepass', 'RC number lookup, challan status'],
        ['Rapidor', 'RC lookup fallback provider'],
        ['NHTSA', 'VIN number decoding'],
        ['Cyepro DMS', 'Dealer management system inventory sync'],
        ['Claude AI (Anthropic)', 'Vehicle description generation'],
        ['Cloudflare', 'DNS management, domain verification'],
        ['GoDaddy', 'Domain registration'],
        ['Vercel', 'Hosting, deployment, edge functions'],
        ['Upstash Redis', 'Domain caching, rate limiting'],
        ['Google Places', 'Review sync from Google Maps'],
        ['Meta Graph API', 'Facebook & Instagram posting'],
        ['Twitter/X API', 'Twitter posting'],
        ['Resend', 'Transactional emails'],
        ['Sentry', 'Error monitoring'],
        ['Google Maps', 'Location embedding'],
    ]
)

add_line()

# --- API Statistics ---
add_heading('API Statistics')
add_bullet('85+ API endpoints')
add_bullet('3 vehicle categories with dedicated APIs each')
add_bullet('12+ external service integrations')
add_bullet('Rate limiting on all critical endpoints')
add_bullet('Redis caching for performance')
add_bullet('Idempotency on payment operations')
add_bullet('Credit tracking for paid API calls')
add_bullet('JWT authentication on all protected routes')
add_bullet('Row-level security \u2014 dealers only access their own data')

add_line()

# --- Payment & Pricing Model ---
add_heading('Payment & Pricing Model')
add_bullet('Dealers can start with a free plan (subdomain website with core features)')
add_bullet('Upgrade to Pro/Premium/Enterprise for custom domains, advanced analytics, and priority support')
add_bullet('RC lookups charged at \u20b93 per call with usage dashboard')
add_bullet('Vehicle booking payments collected online')
add_bullet('All payments processed securely through Razorpay with webhook-based lifecycle management')

add_line()

# --- Why Dealers Should Use This ---
add_heading('Why Dealers Should Use This')
add_table(
    ['Without Dealer Site Pro', 'With Dealer Site Pro'],
    [
        ['No website or an outdated one', 'Professional site live in minutes'],
        ['Leads tracked on paper/WhatsApp', 'Full CRM with priority & status tracking'],
        ['No online inventory', 'Searchable catalog with filters & photos'],
        ['Customers can\u2019t find you on Google', 'SEO-optimized pages that rank locally'],
        ['Pay agencies lakhs for a basic site', 'Start free, upgrade when ready'],
        ['Separate tools for everything', 'One dashboard for inventory, leads, reviews, analytics'],
        ['No customer engagement tools', 'EMI calculator, test drive booking, push notifications'],
        ['Manual vehicle data entry', 'RC lookup auto-fills details in seconds'],
        ['Can\u2019t verify vehicle history', 'Surepass shows challans, insurance, fitness status'],
        ['No social media presence', 'Auto-post listings to Facebook, Instagram, Twitter'],
        ['No analytics', 'Track views, leads, conversions, traffic sources'],
    ]
)

add_line()

# --- Technical Highlights ---
add_heading('Key Technical Highlights (For Technical Stakeholders)')
add_bullet('Built on modern web stack \u2014 Next.js, Supabase, TypeScript')
add_bullet('Multi-tenant architecture \u2014 one platform serves unlimited dealers')
add_bullet('Server-rendered pages for fast loading and SEO')
add_bullet('Edge caching for instant domain resolution')
add_bullet('Row-level security \u2014 each dealer only sees their own data')
add_bullet('Integrates with Surepass (RC verification), Razorpay (payments), Cloudflare (DNS), and Cyepro (DMS)')
add_bullet('40+ database tables, 85+ API endpoints, 142 UI components')

add_line()

# --- In One Line ---
add_heading('In One Line')
p = doc.add_paragraph()
run = p.add_run('Dealer Site Pro is the Shopify for Indian car, bike, and auto dealerships \u2014 everything a dealer needs to get online, manage inventory, capture leads, and grow their business, all from one platform.')
run.bold = True
run.font.size = Pt(12)

# Save
output_path = r"C:\Users\Ravi Abhinav\Desktop\dealer-site-pro\docs\Dealer-Site-Pro-Client-Overview.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
